from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def add_heading_with_number(document, text, level):
    """Add a numbered heading to the document."""
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(14 if level == 1 else 12)
    return paragraph

def add_subheading(document, text):
    """Add a subheading to the document."""
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    return paragraph

def add_content(document, text):
    """Add content text to the document."""
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    run.font.size = Pt(10)
    return paragraph

def create_manual():
    # Create a new Document
    doc = Document()
    
    # Set document title
    title = doc.add_heading('NETWORK VULNERABILITY REPORT GENERATOR', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add subtitle
    subtitle = doc.add_paragraph('ADVANCED USER MANUAL')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add Table of Contents
    doc.add_heading('Table of Contents', 1)
    
    # Introduction
    doc.add_heading('1. Introduction', 1)
    add_subheading(doc, 'Overview')
    add_content(doc, 'The Network Vulnerability Report Generator is a robust and intelligent tool designed to analyze network security, detect vulnerabilities, and generate detailed PDF reports. It serves as a crucial resource for network administrators, cybersecurity professionals, and IT enthusiasts seeking to strengthen their network security posture.')
    
    add_subheading(doc, 'Key Features')
    features = [
        'Automated Network Scanning: Discovers devices connected to the local network.',
        'Vulnerability Assessment: Identifies potential security risks and weaknesses.',
        'Manufacturer Identification: Retrieves device details using a MAC address database.',
        'Customizable PDF Reports: Allows selection of report components for tailored insights.',
        'User-Friendly Interface: Simplifies complex network security tasks.',
        'Cross-Platform Compatibility: Supports major operating systems with Python.'
    ]
    for feature in features:
        add_content(doc, feature)
    
    # System Requirements
    doc.add_heading('2. System Requirements', 1)
    
    add_subheading(doc, 'Operating System')
    add_content(doc, '• Windows 10/11 (Recommended)\n• Linux (Ubuntu 20.04+)\n• macOS 11+')
    
    add_subheading(doc, 'Hardware Requirements')
    add_content(doc, '• Minimum: 4GB RAM, Dual-Core Processor\n• Recommended: 8GB RAM, Quad-Core Processor')
    
    add_subheading(doc, 'Software Dependencies')
    deps = [
        'Python 3.7 or higher',
        'Required Python Libraries:',
        '• tkinter (GUI support)',
        '• psutil (System monitoring)',
        '• scapy (Network packet manipulation)',
        '• reportlab (PDF generation)',
        '• matplotlib (Graph plotting)',
        '• ipaddress (Network address management)'
    ]
    for dep in deps:
        add_content(doc, dep)
    
    # Installation & Setup
    doc.add_heading('3. Installation & Setup', 1)
    
    add_subheading(doc, 'Step 1: Install Python & Dependencies')
    add_content(doc, 'Ensure Python is installed. Then, execute:\npip install -r requirements.txt')
    
    add_subheading(doc, 'Step 2: Running the Application')
    add_content(doc, 'Launch the program using:\npython main.py\n\nFor Linux/macOS users:\nsudo python3 main.py\n(Admin privileges may be required for full network scanning.)')
    
    # Application Workflow
    doc.add_heading('4. Application Workflow', 1)
    workflow = [
        '1. Select Network Interface: Choose the appropriate network adapter.',
        '2. Configure Report Settings: Enable/disable specific report components.',
        '3. Scan Network: Application performs real-time device discovery and vulnerability analysis.',
        '4. Generate Report: PDF report is compiled with identified risks and device information.',
        '5. Export & Review: Save and analyze the generated report for network security insights.'
    ]
    for step in workflow:
        add_content(doc, step)
    
    # User Interface Overview
    doc.add_heading('5. User Interface Overview', 1)
    add_subheading(doc, 'Main Dashboard Elements')
    elements = [
        '1. Network Interface Selector: Dropdown menu to select active network adapter.',
        '2. Report Configuration Panel: Customizable options for included report details.',
        '3. Status Log: Displays real-time scanning progress.',
        '4. Progress Indicator: Visualizes scanning and report generation.',
        '5. Export PDF Button: Generates and saves the report.'
    ]
    for element in elements:
        add_content(doc, element)
    
    # Network Scanning & Interface Selection
    doc.add_heading('6. Network Scanning & Interface Selection', 1)
    
    add_subheading(doc, 'Selecting a Network Interface')
    interface_steps = [
        '1. Open the application and locate the Interface Selector.',
        '2. Choose the appropriate network adapter (e.g., Ethernet, WiFi).',
        '3. The tool will validate the selection and display network details.'
    ]
    for step in interface_steps:
        add_content(doc, step)
    
    add_subheading(doc, 'Performing a Network Scan')
    scan_steps = [
        '1. Click "Start Scan" to initiate the device discovery process.',
        '2. The application will map network topology and assess vulnerabilities.',
        '3. Scanning progress will be displayed in the status log.'
    ]
    for step in scan_steps:
        add_content(doc, step)
    
    # Customizing Report Parameters
    doc.add_heading('7. Customizing Report Parameters', 1)
    add_subheading(doc, 'Options')
    options = [
        'Users can customize the report with the following options:',
        '1. Vulnerability Report (Includes identified threats and security recommendations)',
        '2. Device Inventory (Lists discovered devices with MAC addresses, manufacturers, and hostnames)',
        '3. Network Statistics (Graphs and data insights on network traffic and device activity)',
        '4. Security Risk Levels (Categorization of risks as Critical, High, Medium, Low)'
    ]
    for option in options:
        add_content(doc, option)
    
    # Generating & Exporting Reports
    doc.add_heading('8. Generating & Exporting Reports', 1)
    
    add_subheading(doc, 'Steps')
    report_steps = [
        '1. Select a network interface.',
        '2. Configure report options.',
        '3. Click "Generate Report".',
        '4. Choose a save location and file name.',
        '5. The PDF report will be generated and saved automatically.'
    ]
    for step in report_steps:
        add_content(doc, step)
    
    add_subheading(doc, 'Supported Export Formats')
    formats = [
        '1. PDF (Default, recommended for security analysis)',
        '2. CSV (Optional for structured data processing)'
    ]
    for format in formats:
        add_content(doc, format)
    
    # Understanding the PDF Report
    doc.add_heading('9. Understanding the PDF Report', 1)
    add_subheading(doc, 'Report Sections')
    sections = [
        '1. Title Page: Report title, timestamp, and author details.',
        '2. System Overview: Host device specifications, OS details.',
        '3. Network Summary: Interface details, IP configuration.',
        '4. Device Inventory: List of detected devices with details.',
        '5. Vulnerability Assessment: Security risks categorized by severity.',
        '6. Recommendations: Best practices for improving security.',
        '7. Network Topology Map: Visual representation of network connections.'
    ]
    for section in sections:
        add_content(doc, section)
    
    # Security & Best Practices
    doc.add_heading('10. Security & Best Practices', 1)
    add_subheading(doc, 'Guidelines')
    guidelines = [
        '1. Run the tool with administrator privileges for best results.',
        '2. Ensure firewall rules do not block ARP scanning.',
        '3. Store reports securely to prevent unauthorized access.',
        '4. Regularly update Python dependencies and MAC database.',
        '5. Use in compliance with ethical hacking laws and regulations.'
    ]
    for guideline in guidelines:
        add_content(doc, guideline)
    
    # Troubleshooting & Error Handling
    doc.add_heading('11. Troubleshooting & Error Handling', 1)
    add_subheading(doc, 'Common Issues & Solutions')
    
    # Create a table for issues
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    
    # Add header row
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Issue'
    header_cells[1].text = 'Cause'
    header_cells[2].text = 'Solution'
    
    # Add data rows
    issues = [
        ('No devices found', 'Network adapter issue', 'Ensure the correct interface is selected'),
        ('PDF generation fails', 'File permission issue', 'Run as administrator or change save location'),
        ('Application crashes', 'Missing dependencies', 'Reinstall required Python libraries'),
        ('Incorrect device info', 'MAC database outdated', 'Update OUI database (oui.csv)')
    ]
    
    for issue in issues:
        row_cells = table.add_row().cells
        row_cells[0].text = issue[0]
        row_cells[1].text = issue[1]
        row_cells[2].text = issue[2]
    
    # Frequently Asked Questions
    doc.add_heading('12. Frequently Asked Questions', 1)
    add_subheading(doc, 'General')
    faqs = [
        'Q: Why does the tool require administrator privileges?',
        'A: Some network scanning operations (e.g., ARP scans) require elevated access.',
        '',
        'Q: How frequently should I run vulnerability scans?',
        'A: At least once a month or after significant network changes.',
        '',
        'Q: Can the tool scan remote networks?',
        'A: No, it only scans locally connected networks.',
        '',
        'Q: Is report customization possible?',
        'A: Currently, users can select different sections, but full customization features are planned for future updates.'
    ]
    for faq in faqs:
        add_content(doc, faq)
    
    # Contact & Support
    doc.add_heading('13. Contact & Support', 1)
    add_subheading(doc, 'Information')
    contact_info = [
        'For additional support or bug reports, contact:',
        '📧 Support Email: support@networkvulnscanner.com',
        '🌐 Official Website: www.networkvulnscanner.com',
        '📖 Documentation: docs.networkvulnscanner.com'
    ]
    for info in contact_info:
        add_content(doc, info)
    
    # Copyright
    doc.add_heading('Copyright', 1)
    add_content(doc, '© 2025 Network Vulnerability Report Generator - All Rights Reserved.')
    
    # Save the document
    doc.save('Network_Vulnerability_Report_Generator_Manual.docx')

if __name__ == '__main__':
    create_manual() 